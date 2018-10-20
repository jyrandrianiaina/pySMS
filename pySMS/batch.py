from docx import Document
from docx.shared import Inches
import datetime
import pyping
import smtplib
from eml.MIMEMultipart import MIMEMultipart
from eml.MIMEText import MIMEText
from eml.MIMEBase import MIMEBase
from eml import encoders


def creating_smokeping_checkresult():
    recordset = [ ]
    document = Document()
    today = datetime.datetime.now()
    document.add_heading('Automate check process', 0)
    p = document.add_paragraph('This is an automatre process')
    p.add_run().add_break()
    p.add_run('datetime: %s' %today)

    document.add_heading('Tittle', level=1)
    document.add_paragraph('Tittle', style='ListNumber')
    

    document.add_heading('DNS', level=2)
    document.add_paragraph('Google DNS 1', style='ListNumber')
    document.add_picture('capture.png', width=Inches(5))
    document.add_paragraph('Google DN2', style='ListNumber')
    document.add_picture('capture2.png', width=Inches(5))

    table = document.add_table(rows=1, cols=3)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'IP'
    hdr_cells[1].text = 'Disponibilite'
    hdr_cells[2].text = 'Min/Max/Moyenne TTR en ms'
    my_server = ["8.8.8.8","8.8.4.4","208.67.222.222"]
    for myip in my_server:
        response = pyping.ping(myip,timeout=10000)
        if response.ret_code == 0:
            avlability = "Oui"
            if myip == "8.8.8.8":
                myip = "Google DNS 1/ 8.8.8.8"
            elif myip == "8.8.4.4":
               myip = "Google DNS 2/ 8.8.4.4"
            elif myip == "208.67.222.222":
               myip = "OpenDNS/ 208.67.222.222"
        else:
            avlability = "Non"
            response.max_rtt = " "
            response.avg_rtt = " "
        recordset.extend([[myip,avlability,response.max_rtt+ ' / '+response.avg_rtt]])

    for item in recordset:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item[0])
        row_cells[1].text = str(item[1])
        row_cells[2].text = item[2]
    document.save("allresponse.docx")

creating_smokeping_checkresult()
